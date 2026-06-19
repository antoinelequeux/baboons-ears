#!/usr/bin/env python3

import json
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
DATA_PATH = ROOT / "data" / "individuals.json"
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "Baboons_Ears_User_Guide.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(18)
        paragraph.space_before = Pt(12)
        paragraph.space_after = Pt(6)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(14)
        paragraph.space_before = Pt(10)
        paragraph.space_after = Pt(4)
    else:
        run.bold = True
        run.font.size = Pt(11)
        paragraph.space_before = Pt(6)
        paragraph.space_after = Pt(2)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        run.font.size = Pt(10.5)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        run.font.size = Pt(10.5)


def load_cards():
    payload = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    return payload, payload["cards"]


def build_document():
    payload, cards = load_cards()
    DOCS_DIR.mkdir(exist_ok=True)

    troop_counts = Counter(card["troop"] for card in cards)
    age_sex_counts = Counter(card["ageSex"] for card in cards)
    age_group_counts = Counter(card["ageGroup"] for card in cards)
    single_image_cards = [card for card in cards if len(card["images"]) == 1]
    unique_mark_codes = sorted({card["marks"]["right"] for card in cards} | {card["marks"]["left"] for card in cards})

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Baboons Ears\nComplete User Guide")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Offline field-training manual for the baboon ear identification PWA\n"
        f"Generated on {date.today().isoformat()}"
    )
    run.font.size = Pt(11)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run(
        "This guide explains what is inside the app, how the study modes work, "
        "how the mark codes are displayed, how to install the app on phone or computer, "
        "and how to keep it available without internet."
    )
    run.font.size = Pt(10.5)

    document.add_page_break()

    add_heading(document, "1. What This App Is", level=1)
    add_body(
        document,
        "Baboons Ears is a lightweight progressive web app (PWA) designed for baboon ear-mark training. "
        "It lets a user browse known individuals, study them as flashcards, and keep simple training statistics. "
        "The app is built to keep working offline after the first successful online load.",
    )
    add_bullets(
        document,
        [
            "Browse all individuals as visual cards with ear photos, troop, age/sex, and mark notation.",
            "Study in Normal mode: ears to name.",
            "Study in Hard mode: mark notation to name.",
            "Answer either with a four-choice quiz or by typing the baboon name.",
            "Filter study sessions by troop, age group, and session size.",
            "Track overall accuracy and review the most recent sessions.",
            "Run offline once the app shell, dataset, and images have been cached on the device.",
        ],
    )

    add_heading(document, "2. What Is Included in the Current Dataset", level=1)
    add_body(
        document,
        "The current app dataset comes from the exported training material and a supplementary PDF-based image extraction pass. "
        "The figures below reflect the files currently present in the project.",
    )

    stats_table = document.add_table(rows=1, cols=2)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_table.style = "Table Grid"
    set_cell_text(stats_table.rows[0].cells[0], "Dataset item", bold=True)
    set_cell_text(stats_table.rows[0].cells[1], "Value", bold=True)
    for cell in stats_table.rows[0].cells:
        shade_cell(cell, "D9EAD3")

    stats_rows = [
        ("Deck name", payload["deckName"]),
        ("Total cards", str(payload["cardCount"])),
        ("Troop J cards", str(troop_counts.get("J", 0))),
        ("Troop L cards", str(troop_counts.get("L", 0))),
        ("Adult cards", str(age_group_counts.get("adult", 0))),
        ("Juvenile cards", str(age_group_counts.get("juvenile", 0))),
        ("Subadult cards", str(age_group_counts.get("subadult", 0))),
        ("Cards with no image", "0"),
        ("Cards with one visible ear image only", str(len(single_image_cards))),
    ]
    for label, value in stats_rows:
        cells = stats_table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], value)

    add_body(
        document,
        "Age/sex categories currently present in the app are Adult Female, Adult Male, Juvenile Female, Juvenile Male, and Subadult Male.",
    )
    age_table = document.add_table(rows=1, cols=2)
    age_table.style = "Table Grid"
    age_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(age_table.rows[0].cells[0], "Age / sex", bold=True)
    set_cell_text(age_table.rows[0].cells[1], "Cards", bold=True)
    for cell in age_table.rows[0].cells:
        shade_cell(cell, "D9EAD3")
    for label, count in sorted(age_sex_counts.items()):
        cells = age_table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], str(count))

    if single_image_cards:
        add_body(
            document,
            "Known partial record: Sarcomainf24 currently has one visible ear image in the project because only one clear ear image was available from the source material used for supplementation.",
        )

    add_heading(document, "3. Main Sections of the App", level=1)
    add_bullets(
        document,
        [
            "Browse: a searchable visual catalogue of all baboons in the dataset.",
            "Flashcards: the active training section where study sessions happen.",
            "Field Guide: a short in-app reminder about offline installation and use.",
            "Progress: a summary of total accuracy plus the most recent training sessions.",
        ],
    )

    add_heading(document, "4. Browse Mode", level=1)
    add_body(
        document,
        "Browse mode is the reference view. It is useful when someone wants to inspect individuals freely before starting a timed or structured study session.",
    )
    add_bullets(
        document,
        [
            "Search by name, troop, age/sex, or mark text.",
            "Filter by troop.",
            "Filter by age/sex.",
            "See each card with one or two ear images, the baboon name, troop, age/sex, and mark label.",
            "Clear all browse filters with one button.",
        ],
    )

    add_heading(document, "5. Flashcard Study Modes", level=1)
    add_body(
        document,
        "The study section uses a flashcard logic rather than a classic linear quiz. Each card presents one prompt, reveals the correct answer, then lets the user move to the next card.",
    )
    add_bullets(
        document,
        [
            "Normal (ears to name): the front of the card shows the ear photos and the user must identify the baboon.",
            "Hard (marks to name): the front of the card shows only the mark notation, for example 'R: TB • L: MB', and the user must identify the baboon from the notation alone.",
            "Quiz (4 choices): the app proposes four names and highlights green for the correct answer and red for a wrong selected answer.",
            "Write the name: the user types the answer manually and the app checks it against the stored baboon name.",
        ],
    )

    add_heading(document, "6. Study Filters and Session Controls", level=1)
    add_bullets(
        document,
        [
            "Troop filter: All troops, J only, or L only.",
            "Age-group filter: All ages, Adult only, Juvenile only, or Subadult only.",
            "Session size: 10 cards, 20 cards, 50 cards, or All cards.",
            "Start session: builds a new queue from the currently allowed study pool.",
        ],
    )
    add_body(
        document,
        "A session stops when the chosen target size has been reached. If the user selects 10 cards, the app records a 10-card session rather than continuing indefinitely.",
    )

    add_heading(document, "7. Intelligent Card Rotation", level=1)
    add_body(
        document,
        "The app uses a simple weighted system so difficult baboons return more often. This is not random repetition only: it is designed to bring back individuals that were missed more often or that do not yet have a good success streak.",
    )
    add_bullets(
        document,
        [
            "Each card stores how many times it has been seen, answered correctly, answered wrongly, and how long the current success streak is.",
            "Wrong answers increase a card's future weight strongly.",
            "Short streaks also increase the chance of coming back soon.",
            "The result is a practical field-training loop where weak points are reviewed more often than easy individuals.",
        ],
    )

    add_heading(document, "8. Progress Tracking", level=1)
    add_body(
        document,
        "Progress is stored locally on the device through browser local storage. It stays on that same device unless the browser data is cleared.",
    )
    add_bullets(
        document,
        [
            "Overall accuracy percentage across completed sessions.",
            "Total number of correct answers out of total reviewed cards.",
            "A visual progress bar.",
            "A recent-session history showing mode, answer style, score, session size, troop filter, and age filter.",
            "Per-card memory used by the repetition logic.",
        ],
    )

    add_heading(document, "9. Understanding the Ear-Mark Notation", level=1)
    add_body(
        document,
        "The app preserves the original mark notation from the source deck. It does not invent or reinterpret the ear-mark strings. This is important because the app is meant to train the user on the exact codes already used in the field material.",
    )
    add_bullets(
        document,
        [
            "R means right ear.",
            "L means left ear.",
            "A label such as 'R: TB • L: MB' means the right-ear code is TB and the left-ear code is MB.",
            "A value of '0' means no specific code was recorded for that ear in the source material.",
            "Combined strings such as TB, TMB, TTMB, sT, or sTB are displayed exactly as they appear in the exported dataset.",
            "In Hard mode, these notation strings are the prompt the user studies from.",
        ],
    )
    add_body(
        document,
        "Important caution: the app can document and display the code vocabulary, but the exact biological or field-definition expansion of each letter string should be confirmed against the project's official ear-notch legend if you need a formal field protocol. In other words, the app is faithful to the source codes, but it is not itself the authoritative decoding key for the letters.",
    )

    add_heading(document, "10. Mark Code Vocabulary Seen in the Current Dataset", level=1)
    add_body(
        document,
        "The following ear-code values are currently present in the dataset. They are listed exactly as stored in the source-derived JSON.",
    )
    codes_table = document.add_table(rows=1, cols=3)
    codes_table.style = "Table Grid"
    codes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = codes_table.rows[0].cells
    set_cell_text(header_cells[0], "Code", bold=True)
    set_cell_text(header_cells[1], "How to read it in the app", bold=True)
    set_cell_text(header_cells[2], "Interpretation status", bold=True)
    for cell in header_cells:
        shade_cell(cell, "D9EAD3")
    for code in unique_mark_codes:
        cells = codes_table.add_row().cells
        set_cell_text(cells[0], code)
        set_cell_text(cells[1], f"Shown literally as '{code}' on the relevant ear")
        if code == "0":
            set_cell_text(cells[2], "No recorded code for that ear")
        else:
            set_cell_text(cells[2], "Source code preserved; confirm official legend for exact meaning")

    add_heading(document, "11. Example Study Prompts", level=1)
    add_bullets(
        document,
        [
            "Normal example: two ear photos are shown; the user must identify the baboon's name.",
            "Hard example: 'R: TBB • L: T' appears; the user must recall the correct name from the notation.",
            "Quiz example: four names are offered; selecting the correct one turns the right choice green.",
            "Typed-answer example: the user enters a name manually and the app compares it case-insensitively.",
        ],
    )

    add_heading(document, "12. How to Install on iPhone", level=1)
    add_numbered(
        document,
        [
            "Deploy the project to a web host such as GitHub Pages.",
            "Open the site in Safari while the phone is online.",
            "Wait for the first full load so the app shell, JSON, and images can cache.",
            "Tap Share, then choose Add to Home Screen.",
            "Open the app once from the new icon while still online.",
            "Test it with Airplane Mode enabled before relying on it in the field.",
        ],
    )
    add_body(
        document,
        "Important iPhone note: iOS can be more aggressive than Android about clearing cached website data if storage is low or if the app is unused for a long time. Always test the installed app on the real device shortly before field use.",
    )

    add_heading(document, "13. How to Install on Android", level=1)
    add_numbered(
        document,
        [
            "Deploy the project to a web host such as GitHub Pages.",
            "Open the site in Chrome or another Chromium-based browser while online.",
            "Wait for the first full load.",
            "Choose Add to Home Screen or Install App when offered by the browser.",
            "Launch it once from the installed icon before going offline.",
            "Test it with Wi-Fi and mobile data both turned off.",
        ],
    )
    add_body(
        document,
        "Android is generally more permissive and reliable for offline PWAs, especially when the image cache has been warmed successfully on first use.",
    )

    add_heading(document, "14. How to Use It on a Computer", level=1)
    add_bullets(
        document,
        [
            "Simplest option: use the deployed website in a desktop browser.",
            "For Chrome or Edge, you can usually install it as a desktop app from the browser menu.",
            "If you are working locally, open it through a local web server rather than file:// because service workers need an HTTP context.",
            "A simple local server can be started from the project folder with 'python3 -m http.server 8000', then opened at http://localhost:8000.",
        ],
    )
    add_body(
        document,
        "On desktop, offline behavior still depends on first loading the app successfully so the service worker and cached assets are available.",
    )

    add_heading(document, "15. How Offline Mode Works", level=1)
    add_body(
        document,
        "This project is built as a progressive web app. A service worker caches the app shell files, the JSON dataset, and the ear photos. After that first successful caching step, the app can reopen without network access.",
    )
    add_bullets(
        document,
        [
            "Core cached files include the HTML, CSS, JavaScript, manifest, service worker, dataset JSON, and icon files.",
            "Image warm-up logic also adds ear images to the browser cache.",
            "The connection indicator changes between an online/cached state and an offline state.",
            "If the device has never finished the first online load, offline use may be incomplete.",
        ],
    )

    add_heading(document, "16. How to Update the Website After Changes", level=1)
    add_numbered(
        document,
        [
            "Edit the project files locally.",
            "Commit and push the changes to the GitHub repository.",
            "Wait for GitHub Pages to redeploy.",
            "Refresh the live website on the device.",
            "If an old version seems stuck, close the installed app fully and reopen it.",
            "If needed, remove the home-screen app and install it again to force a clean refresh.",
        ],
    )
    add_body(
        document,
        "Because PWAs use cached assets, a browser can occasionally keep an older version longer than expected. Reopening after the service-worker cache version changes usually resolves that.",
    )

    add_heading(document, "17. Known Limits and Good Practice", level=1)
    add_bullets(
        document,
        [
            "The app stores training progress locally on one device; it is not a cloud-synced score system.",
            "If browser data is cleared, the local progress history is lost.",
            "The exact semantic decoding of every mark string should be validated against the team's official notch legend.",
            "One record currently has only one visible ear image in the project files.",
            "Before field deployment, always test on the real phone in offline mode.",
        ],
    )

    document.add_page_break()
    add_heading(document, "18. Quick Glossary", level=1)
    glossary_rows = [
        ("Browse", "The visual catalogue of all cards."),
        ("Flashcard", "A study card with a front prompt and a revealed answer."),
        ("Normal mode", "Ears to name."),
        ("Hard mode", "Mark notation to name."),
        ("Quiz mode", "Four answer choices."),
        ("Write the name", "Free-text answer entry."),
        ("Troop", "Group label currently shown as J or L."),
        ("Age / Sex", "Detailed category such as Adult Female or Juvenile Male."),
        ("Age group", "Broader filter: adult, juvenile, or subadult."),
        ("Session size", "How many cards are included in one study run."),
        ("Offline or cached", "The app has enough local files stored to run without network access."),
    ]
    glossary = document.add_table(rows=1, cols=2)
    glossary.style = "Table Grid"
    glossary.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(glossary.rows[0].cells[0], "Term", bold=True)
    set_cell_text(glossary.rows[0].cells[1], "Meaning", bold=True)
    for cell in glossary.rows[0].cells:
        shade_cell(cell, "D9EAD3")
    for term, meaning in glossary_rows:
        cells = glossary.add_row().cells
        set_cell_text(cells[0], term)
        set_cell_text(cells[1], meaning)

    add_heading(document, "19. File Structure in the Project", level=1)
    add_bullets(
        document,
        [
            "index.html: main page structure.",
            "style.css: layout and visual styling.",
            "app.js: data loading, browse filters, flashcard logic, progress tracking, and cache warm-up.",
            "manifest.json: install metadata for the PWA.",
            "service-worker.js: offline caching logic.",
            "data/individuals.json: generated dataset with names, groups, marks, and image paths.",
            "images/: ear-photo assets used by the cards.",
            "scripts/: helper scripts for rebuilding data and generating project outputs.",
        ],
    )

    add_heading(document, "20. Final Summary", level=1)
    add_body(
        document,
        "In practice, this app is a focused offline study tool for baboon identification. It combines ear photographs, source mark notation, flashcard modes, device-local progress, and PWA offline delivery. Once deployed and installed correctly, it provides a cleaner field-training workflow than a generic quiz tool because it respects the troop filters, age-group filters, ear-mark prompts, and repeated review of difficult individuals.",
    )

    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_document()
    print(path)
